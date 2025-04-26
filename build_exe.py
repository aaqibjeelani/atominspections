#!/usr/bin/env python3
"""
Build script to create a standalone executable for Word Document Merger
"""
import os
import sys
import subprocess
import shutil
from create_icon import create_simple_icon

def install_requirements():
    """Install required packages for building the executable"""
    packages = [
        'pyinstaller',
        'pillow'  # For icon creation
    ]
    
    print("Installing required packages...")
    subprocess.call([sys.executable, '-m', 'pip', 'install'] + packages)
    
def create_executable():
    """Create the standalone executable using PyInstaller"""
    # First, create the icon
    print("Creating application icon...")
    icon_created = create_simple_icon()
    icon_path = 'app_icon.ico' if icon_created else 'app_icon.png'
    
    # Create a directory for the sample template if it doesn't exist
    if not os.path.exists('sample'):
        os.makedirs('sample')
        print("Created 'sample' directory for template files")
    
    # Check if template exists, if not create a simple one
    if not os.path.exists('sample/template.docx'):
        try:
            from docx import Document
            print("Creating a simple template document...")
            doc = Document()
            doc.add_heading('Word Document Merger Template', 0)
            doc.add_paragraph('This is a simple template document.')
            table = doc.add_table(rows=1, cols=7)
            header_cells = table.rows[0].cells
            header_cells[0].text = 'REF'
            header_cells[1].text = 'Item 1'
            header_cells[2].text = 'Item 2'
            header_cells[3].text = 'Item 3'
            header_cells[4].text = 'Item 4'
            header_cells[5].text = 'Item 5'
            header_cells[6].text = 'Image'
            doc.save('sample/template.docx')
            print("Template created: sample/template.docx")
        except Exception as e:
            print(f"Could not create template: {e}")
            print("Please provide your own template.docx in the sample directory.")
    
    # Build the executable with PyInstaller
    print("\nBuilding executable with PyInstaller...")
    pyinstaller_cmd = [
        'pyinstaller',
        '--name=WordDocumentMerger',
        '--onefile',  # Create a single executable file
        f'--icon={icon_path}',
        '--noconsole',  # Don't show console window
        '--add-data=sample;sample',  # Include sample directory
        'merge_word_docs.py'
    ]
    
    # Add the icon file to be included
    if os.path.exists(icon_path):
        pyinstaller_cmd.append(f'--add-data={icon_path};.')
    
    subprocess.call(pyinstaller_cmd)
    
    print("\nBuild completed!")
    print("The executable can be found in the 'dist' directory")
    print("Note: You need to include the 'sample' directory with the executable for the template.")

if __name__ == "__main__":
    install_requirements()
    create_executable()
    
    # Copy sample directory to dist for convenience
    if os.path.exists('dist') and os.path.exists('sample'):
        if not os.path.exists('dist/sample'):
            shutil.copytree('sample', 'dist/sample')
            print("Copied sample directory to dist")
    
    print("\nDone! Your standalone executable has been created.")
    print("Developed by Aaqib Jeelani - instagram.com/aaqibjeelani") 