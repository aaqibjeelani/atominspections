#!/usr/bin/env python3
import os
import sys
import tempfile
import shutil
import subprocess
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk

# Fixed template path
TEMPLATE_PATH = "sample/template.docx"

def is_header_row(row_index):
    """
    Determine if a row is a header row. Header rows include:
    - The first row (index 0)
    - Every 4th row after that (rows 5, 9, 13, etc. which have indices 4, 8, 12, etc.)
    """
    # First row is always a header
    if row_index == 0:
        return True
    
    # Every 4th row after row 1 is a header (row 5, 9, 13, etc.)
    # Using modulo 4 with offset to identify these rows
    return (row_index % 4 == 0) and row_index > 0

def file_exists(file_path):
    """Check if a file exists and is accessible"""
    return os.path.isfile(file_path) and os.access(file_path, os.R_OK)

def add_ref_numbers_with_word(doc_path, output_path, start_number):
    """Process a document to add REF numbers to tables using Word COM automation"""
    try:
        import win32com.client
        import pythoncom
        
        # Initialize COM
        pythoncom.CoInitialize()
        
        # Create Word application
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        
        try:
            # Validate file exists
            if not file_exists(doc_path):
                print(f"Error: File not found or not accessible: {doc_path}")
                raise FileNotFoundError(f"File not found: {doc_path}")
            
            print(f"Opening document with Word: {doc_path}")
            # Get absolute path
            abs_path = os.path.abspath(doc_path)
            
            # Open the document
            doc = word.Documents.Open(abs_path)
            
            # Start with provided serial number
            serial_number = start_number
            table_count = 0
            
            print(f"Document has {doc.Tables.Count} tables")
            
            # Process all tables in the document
            for i in range(1, doc.Tables.Count + 1):
                table = doc.Tables(i)
                table_count += 1
                row_count = table.Rows.Count
                print(f"Processing table {i} with {row_count} rows")
                
                # Check if the table has columns
                if table.Columns.Count < 1:
                    print(f"  Table {i} has no columns, skipping")
                    continue
                
                # First scan the table to identify cells with images in the entire table
                cells_with_images = set()
                for j in range(1, row_count + 1):
                    for k in range(1, table.Columns.Count + 1):
                        try:
                            cell = table.Cell(j, k)
                            # Check for inline shapes or regular shapes
                            if cell.Range.InlineShapes.Count > 0 or cell.Range.ShapeRange.Count > 0:
                                print(f"  Detected image in cell at row {j}, column {k}")
                                cells_with_images.add((j, k))
                        except Exception:
                            # If error checking for images, continue
                            pass
                
                # Process each row for adding REF numbers only in first column where appropriate
                for j in range(1, row_count + 1):
                    # Word uses 1-based indexing, convert to 0-based for our header check
                    row_index = j - 1
                    
                    # Skip header rows
                    if is_header_row(row_index):
                        print(f"  Skipping header row {j} (index {row_index})")
                        continue
                    
                    # Skip if the first cell in this row contains an image
                    if (j, 1) in cells_with_images:
                        print(f"  Skipping row {j} - first cell contains an image")
                        serial_number += 1  # Still increment the counter
                        continue
                    
                    # Also skip if any cell in this row contains an image, as it might be a merged row
                    row_has_images = False
                    for k in range(1, table.Columns.Count + 1):
                        if (j, k) in cells_with_images:
                            row_has_images = True
                            break
                    
                    if row_has_images:
                        print(f"  Skipping row {j} - row contains images")
                        serial_number += 1  # Still increment the counter
                        continue
                    
                    # Add REF number to first cell only if no images in row
                    try:
                        # Get first cell in row
                        cell = table.Cell(j, 1)  # 1 is the first column in Word COM
                        
                        # Different approach to set text - more reliable
                        try:
                            # Select the cell
                            cell.Select()
                            # Clear selection and insert text
                            word.Selection.Text = str(serial_number)
                            print(f"  Set REF number {serial_number} in row {j}")
                        except Exception as select_error:
                            print(f"  Error with selection method: {select_error}")
                            # Fallback to range method
                            cell.Range.Text = str(serial_number)
                            print(f"  Set REF number {serial_number} via Range in row {j}")
                            
                        serial_number += 1
                    except Exception as cell_error:
                        print(f"  Error processing cell in row {j}: {cell_error}")
                        traceback.print_exc()
                        serial_number += 1  # Still increment even on error
            
            print(f"Processed {table_count} tables in document")
            
            # Create output directory if it doesn't exist
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir)
            
            # Save the processed document
            abs_output_path = os.path.abspath(output_path)
            print(f"Saving document to: {abs_output_path}")
            doc.SaveAs2(abs_output_path)
            doc.Close(SaveChanges=False)
            
            print(f"Saved processed document to: {output_path}")
            return serial_number
            
        finally:
            # Make sure Word is properly closed
            word.Quit()
            pythoncom.CoUninitialize()
    
    except ImportError:
        print("Warning: win32com not available, falling back to python-docx method")
        return add_ref_numbers(doc_path, output_path, start_number)
    except Exception as e:
        print(f"Error processing document with Word COM: {e}")
        import traceback
        traceback.print_exc()
        # Fall back to python-docx method
        print("Falling back to python-docx method")
        return add_ref_numbers(doc_path, output_path, start_number)

def add_ref_numbers(doc_path, output_path, start_number):
    """Process a document to add REF numbers to tables using python-docx"""
    try:
        # Validate file exists
        if not file_exists(doc_path):
            print(f"Error: File not found or not accessible: {doc_path}")
            raise FileNotFoundError(f"File not found: {doc_path}")
        
        print(f"Opening document with python-docx: {doc_path}")
        # Open the document
        doc = Document(doc_path)
        
        # Start with provided serial number
        serial_number = start_number
        table_count = 0
        
        # Process all tables in the document
        for i, table in enumerate(doc.tables):
            table_count += 1
            print(f"Processing table {i+1} with {len(table.rows)} rows")
            
            # Skip table if it has no columns
            if len(table.columns) < 1:
                print(f"  Table {i+1} has no columns, skipping")
                continue
            
            # First identify all cells containing images
            cells_with_images = set()
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        # Check for drawings or pictures
                        if (para._element.xpath('.//w:drawing') or 
                            para._element.xpath('.//w:pict') or
                            para._element.xpath('.//a:blip')):
                            print(f"  Detected image in cell at row {row_idx+1}, column {col_idx+1}")
                            cells_with_images.add((row_idx, col_idx))
                            break
                
            # Process rows for numbering
            for j, row in enumerate(table.rows):
                # Skip header rows
                if is_header_row(j):
                    print(f"  Skipping header row {j+1}")
                    continue
                
                # Skip if the first cell contains an image
                if (j, 0) in cells_with_images:
                    print(f"  Skipping row {j+1} - first cell contains an image")
                    serial_number += 1  # Still increment counter
                    continue
                
                # Skip if any cell in this row contains an image (may affect layout)
                row_has_images = False
                for col_idx in range(len(row.cells)):
                    if (j, col_idx) in cells_with_images:
                        row_has_images = True
                        break
                
                if row_has_images:
                    print(f"  Skipping row {j+1} - row contains images")
                    serial_number += 1  # Still increment counter
                    continue
                
                # Add REF number to first cell if it doesn't contain images
                if len(row.cells) > 0:
                    first_cell = row.cells[0]
                    
                    # Clear existing content and add REF number
                    if len(first_cell.paragraphs) == 0:
                        # Create paragraph if none exists
                        p = first_cell.add_paragraph()
                        p.text = str(serial_number)
                    else:
                        # Clear existing paragraphs
                        for p in first_cell.paragraphs:
                            for run in p.runs:
                                run.text = ""
                        
                        # Add number to first paragraph
                        if len(first_cell.paragraphs) > 0:
                            first_cell.paragraphs[0].add_run(str(serial_number))
                    
                    print(f"  Set REF number {serial_number} in row {j+1}")
                
                serial_number += 1
        
        print(f"Processed {table_count} tables in document")
        
        # Create output directory if it doesn't exist
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # Save the processed document
        doc.save(output_path)
        print(f"Saved processed document to: {output_path}")
        
        return serial_number
    
    except FileNotFoundError as e:
        print(f"Error: File not found - {e}")
        raise
    except Exception as e:
        print(f"Error processing document: {e}")
        import traceback
        traceback.print_exc()
        raise

def simple_copy_merge(template_path, processed_docs, output_path):
    """A simple fallback method that creates a new document and copies content"""
    try:
        print(f"Starting simple copy merge with {len(processed_docs)} documents")
        
        # Handle missing template
        if not template_path or not file_exists(template_path):
            print("Template not found, using first document as base")
            # Use the first document as our base document if no template
            if processed_docs:
                merged_doc = Document(processed_docs[0])
                print(f"Using {processed_docs[0]} as base document")
                # Process each additional document (skip the first one as it's our base)
                docs_to_process = processed_docs[1:]
            else:
                print("No documents to process")
                return False
        else:
            # Use the template as our base
            print(f"Using template {template_path} as base document")
            merged_doc = Document(template_path)
            # Process all documents
            docs_to_process = processed_docs
            
        # Track if we have added any content
        content_added = False if not docs_to_process else True
            
        # Process each document
        for i, doc_path in enumerate(docs_to_process):
            if not file_exists(doc_path):
                print(f"Warning: File not found, skipping: {doc_path}")
                continue
                
            print(f"Merging content from document: {doc_path}")
            doc = Document(doc_path)
            
            # Add page break before adding new content (except for first document)
            if i > 0 or (template_path and file_exists(template_path)):
                merged_doc.add_page_break()
                print(f"  Added page break before document {i+1}")
            
            # Copy paragraphs
            for para in doc.paragraphs:
                p = merged_doc.add_paragraph()
                for run in para.runs:
                    r = p.add_run(run.text)
                    r.bold = run.bold
                    r.italic = run.italic
                    r.underline = run.underline
            
            # Copy tables - preserving as much formatting as possible
            for table in doc.tables:
                # Create new table with same dimensions
                if len(table.rows) > 0 and len(table.columns) > 0:
                    new_table = merged_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                    
                    # Copy style if available
                    if hasattr(table, 'style') and table.style:
                        try:
                            new_table.style = table.style
                        except Exception as style_error:
                            print(f"  Could not copy table style: {style_error}")
                    
                    # Copy cell contents
                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            if j < len(new_table.rows[i].cells):  # Ensure index exists
                                new_cell = new_table.rows[i].cells[j]
                                
                                # Copy cell content
                                if len(cell.paragraphs) > 0:
                                    # Remove default paragraph if it exists
                                    if len(new_cell.paragraphs) > 0:
                                        new_cell.paragraphs[0].text = ""
                                        
                                    # Copy each paragraph and its runs
                                    for para in cell.paragraphs:
                                        if len(new_cell.paragraphs) > 0 and not new_cell.paragraphs[0].text:
                                            # Use existing paragraph if it's empty
                                            new_para = new_cell.paragraphs[0]
                                        else:
                                            # Otherwise add a new paragraph
                                            new_para = new_cell.add_paragraph()
                                            
                                        # Copy runs with formatting
                                        for run in para.runs:
                                            new_run = new_para.add_run(run.text)
                                            if hasattr(run, 'bold'):
                                                new_run.bold = run.bold
                                            if hasattr(run, 'italic'):
                                                new_run.italic = run.italic
                                            if hasattr(run, 'underline'):
                                                new_run.underline = run.underline
                    
                    # Add paragraph after table (for spacing)
                    merged_doc.add_paragraph()
                    content_added = True
            
        # Only save if we actually added content
        if content_added:
            # Create output directory if it doesn't exist
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir)
                
            # Save merged document
            merged_doc.save(output_path)
            print(f"Merged document saved: {output_path}")
            return True
        else:
            print("No content was added to the merged document")
            return False
        
    except Exception as e:
        print(f"Error in simple document merge: {e}")
        import traceback
        traceback.print_exc()
        return False

def merge_docs_with_msword(output_path, doc_paths, template_path=None):
    """Use Microsoft Word directly via COM objects to merge documents"""
    try:
        import win32com.client
        import pythoncom
        
        # Initialize COM
        pythoncom.CoInitialize()
        
        # Create Word application
        word = win32com.client.Dispatch("Word.Application")
        # For debugging image issues, you might want to see what's happening
        word.Visible = False
        word.DisplayAlerts = False
        
        try:
            # Check if documents exist
            valid_docs = []
            for doc_path in doc_paths:
                if file_exists(doc_path):
                    valid_docs.append(os.path.abspath(doc_path))
                else:
                    print(f"Warning: Document not found, skipping: {doc_path}")
                    
            if not valid_docs:
                print("Error: No valid documents to merge")
                return False
            
            # Create output directory if it doesn't exist
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir)
                
            # For more reliable image handling, use a direct file insertion approach
            # Create a new base document that will contain all others
            print("Creating a base document for insertion...")
            
            # Start with the first document for better preservation of formatting
            # This avoids template formatting issues with images
            if len(valid_docs) > 0:
                print(f"Using first document as base: {valid_docs[0]}")
                shutil.copyfile(valid_docs[0], output_path)
                base_doc = word.Documents.Open(os.path.abspath(output_path))
                
                # Process each additional document (skip the first one as it's our base)
                for i, doc_path in enumerate(valid_docs[1:], 1):
                    print(f"Processing document {i+1}/{len(valid_docs)}: {doc_path}")
                    
                    # Insert a page break at the end before inserting next document
                    base_doc.Content.Collapse(0)  # wdCollapseEnd = 0
                    word.Selection.EndKey(6)  # wdStory = 6 (move to end of document)
                    word.Selection.InsertBreak(7)  # wdPageBreak = 7
                    
                    # Insert the document - this preserves images and formatting better
                    try:
                        # Use InsertFile for direct file insertion
                        word.Selection.InsertFile(doc_path)
                        print(f"  Inserted file directly: {os.path.basename(doc_path)}")
                    except Exception as e:
                        print(f"  Error directly inserting file: {e}")
                        
                        # Fallback to object model copy/paste if direct insertion fails
                        try:
                            print(f"  Trying alternative method...")
                            src_doc = word.Documents.Open(doc_path)
                            src_doc.Content.Copy()
                            src_doc.Close(False)
                            
                            word.Selection.Paste()
                            print(f"  Pasted content from {os.path.basename(doc_path)}")
                        except Exception as paste_error:
                            print(f"  Error with alternative method: {paste_error}")
                
                # Save the merged document
                print(f"Saving merged document to: {output_path}")
                base_doc.Save()
                base_doc.Close(SaveChanges=True)
                
                print(f"Document successfully saved to: {output_path}")
                return True
            else:
                print("No valid documents to process.")
                return False
            
        finally:
            # Make sure Word is properly closed
            word.Quit()
            pythoncom.CoUninitialize()
            
    except ImportError as e:
        print(f"Error importing required modules: {e}")
        return False
    except Exception as e:
        print(f"Error with Word COM: {e}")
        import traceback
        traceback.print_exc()
        return False

def merge_documents(template_path, document_paths, output_path):
    """
    Main merge function that handles the document merging process.
    It first adds REF numbers to each document, then merges them.
    """
    try:
        print(f"Starting merge process...")
        print(f"Using template: {template_path}")
        print(f"Documents to merge: {len(document_paths)}")
        print(f"Output: {output_path}")
        
        # Validate paths
        template_exists = file_exists(template_path)
        if not template_exists:
            print(f"Warning: Template file not found: {template_path}")
            print("Will continue without template")
            
        valid_docs = []
        for doc_path in document_paths:
            if file_exists(doc_path):
                valid_docs.append(doc_path)
            else:
                print(f"Warning: Document not found, will be skipped: {doc_path}")
                
        if not valid_docs:
            print("Error: No valid documents to merge")
            return False
        
        if len(valid_docs) == 1:
            # Special case: just a single document
            print("Only one document to process. Applying REF numbers directly...")
            
            try:
                print(f"Adding REF numbers to single document...")
                # Try to use Word COM automation first (preferred for formatting)
                try:
                    import win32com.client
                    print("Using Word COM for REF numbering...")
                    add_ref_numbers_with_word(valid_docs[0], output_path, 1)
                    print(f"Successfully processed document to: {output_path}")
                    return True
                except (ImportError, Exception) as e:
                    print(f"Word COM method failed: {e}")
                    print("Using python-docx for REF numbering...")
                    add_ref_numbers(valid_docs[0], output_path, 1)
                    print(f"Successfully processed document to: {output_path}")
                    return True
            except Exception as e:
                print(f"Error processing single document: {e}")
                import traceback
                traceback.print_exc()
                return False
            
        # Create a temporary directory for processed files
        temp_dir = tempfile.mkdtemp()
        print(f"Created temporary directory: {temp_dir}")
        
        try:
            # Process each document to add REF numbers
            serial_number = 1
            processed_docs = []
            
            for i, doc_path in enumerate(valid_docs):
                print(f"\nProcessing document {i+1}: {os.path.basename(doc_path)}")
                
                # Create temporary file for processed document
                temp_output = os.path.join(temp_dir, f"doc_{i}.docx")
                
                try:
                    # Try to use Word COM automation first (preferred for formatting)
                    success = False
                    try:
                        import win32com.client
                        print("Using Word COM for REF numbering...")
                        serial_number = add_ref_numbers_with_word(doc_path, temp_output, serial_number)
                        success = True
                    except (ImportError, Exception) as e:
                        print(f"Word COM method failed: {e}")
                        print("Using python-docx for REF numbering...")
                        serial_number = add_ref_numbers(doc_path, temp_output, serial_number)
                        success = True
                        
                    if success:
                        processed_docs.append(temp_output)
                except Exception as e:
                    print(f"Error processing document {doc_path}: {e}")
                    import traceback
                    traceback.print_exc()
                    # If we can't process it for REF numbers, still include the original
                    print(f"Including original document without REF numbers: {doc_path}")
                    processed_docs.append(doc_path)
            
            if not processed_docs:
                print("Error: No documents were successfully processed")
                return False
                
            # Try different merge methods in order of preference
            print("\nAttempting to merge documents...")
            
            # First try using Word COM for direct merging
            try:
                import win32com.client
                print("Using Word COM interface for merging...")
                template_to_use = template_path if template_exists else None
                success = merge_docs_with_msword(output_path, processed_docs, template_to_use)
                if success:
                    print("Merge completed successfully using Word COM!")
                    return True
                else:
                    print("Word COM merge failed. Trying alternative method...")
            except ImportError:
                print("Word COM interface not available. Trying alternative method...")
            except Exception as e:
                print(f"Error with Word COM merge: {e}")
                import traceback
                traceback.print_exc()
                print("Trying alternative method...")
            
            # If COM interface failed, use the simple copy merge approach
            print("Using simple merge method...")
            template_to_use = template_path if template_exists else None
            success = simple_copy_merge(template_to_use, processed_docs, output_path)
            if success:
                print("Merge completed successfully using simple merge!")
                return True
            else:
                # If only one document, just copy it directly
                if len(processed_docs) == 1:
                    print("Only one document. Copying directly...")
                    shutil.copyfile(processed_docs[0], output_path)
                    if os.path.exists(output_path):
                        print(f"Single document copied to: {output_path}")
                        return True
                        
                print("All merge methods failed.")
                return False
                
        finally:
            # Clean up temporary files
            try:
                shutil.rmtree(temp_dir)
                print("Temporary files removed")
            except Exception as e:
                print(f"Warning: Could not remove temporary files: {e}")
    
    except Exception as e:
        print(f"Error in merge process: {e}")
        import traceback
        traceback.print_exc()
        return False

class WordMergerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Word Document Merger")
        self.root.geometry("700x550")  # Slightly taller to accommodate credits
        self.root.minsize(650, 450)
        
        # Add application icon if available
        try:
            self.root.iconbitmap("app_icon.ico")
        except:
            pass  # Icon not found, continue without it
        
        self.create_widgets()
        
    def create_widgets(self):
        # Template path label (read-only)
        template_frame = ttk.LabelFrame(self.root, text="Template Document")
        template_frame.pack(fill="x", expand=False, padx=10, pady=10)
        
        self.template_label = ttk.Label(template_frame, text=f"Using template: {TEMPLATE_PATH}")
        self.template_label.pack(fill="x", expand=True, padx=5, pady=5)
        
        # Frame for documents selection
        docs_frame = ttk.LabelFrame(self.root, text="Documents to Merge")
        docs_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        self.docs_listbox = tk.Listbox(docs_frame, selectmode=tk.SINGLE, height=10)
        self.docs_listbox.pack(side=tk.LEFT, fill="both", expand=True, padx=5, pady=5)
        
        scrollbar = ttk.Scrollbar(docs_frame, orient="vertical", command=self.docs_listbox.yview)
        scrollbar.pack(side=tk.RIGHT, fill="y")
        self.docs_listbox.config(yscrollcommand=scrollbar.set)
        
        # Buttons for document list management
        btn_frame = ttk.Frame(docs_frame)
        btn_frame.pack(fill="x", expand=False, padx=5, pady=5)
        
        self.add_btn = ttk.Button(btn_frame, text="Add Documents", command=self.add_documents)
        self.add_btn.pack(side=tk.LEFT, padx=5)
        
        self.remove_btn = ttk.Button(btn_frame, text="Remove Selected", command=self.remove_document)
        self.remove_btn.pack(side=tk.LEFT, padx=5)
        
        self.clear_btn = ttk.Button(btn_frame, text="Clear All", command=self.clear_documents)
        self.clear_btn.pack(side=tk.LEFT, padx=5)
        
        # Frame for output file
        output_frame = ttk.LabelFrame(self.root, text="Output Document")
        output_frame.pack(fill="x", expand=False, padx=10, pady=10)
        
        self.output_path_var = tk.StringVar()
        self.output_path_entry = ttk.Entry(output_frame, textvariable=self.output_path_var, width=50)
        self.output_path_entry.pack(side=tk.LEFT, fill="x", expand=True, padx=5, pady=5)
        
        self.browse_output_btn = ttk.Button(output_frame, text="Browse...", command=self.browse_output)
        self.browse_output_btn.pack(side=tk.RIGHT, padx=5, pady=5)
        
        # Merge button
        self.merge_btn = ttk.Button(self.root, text="Merge Documents", command=self.merge_documents)
        self.merge_btn.pack(pady=10)
        
        # Status label
        self.status_var = tk.StringVar()
        self.status_label = ttk.Label(self.root, textvariable=self.status_var, wraplength=650)
        self.status_label.pack(fill="x", expand=False, padx=10, pady=5)
        
        # Add credits at the bottom
        credits_label = ttk.Label(self.root, text="Developed by Aaqib Jeelani - instagram.com/aaqibjeelani", 
                                 font=("Arial", 8), foreground="#555555")
        credits_label.pack(side=tk.BOTTOM, pady=10)
        
        # Document paths
        self.document_paths = []
            
    def add_documents(self):
        paths = filedialog.askopenfilenames(
            title="Select Documents to Merge",
            filetypes=[("Word Documents", "*.docx")]
        )
        for path in paths:
            if path not in self.document_paths:
                self.document_paths.append(path)
                self.docs_listbox.insert(tk.END, os.path.basename(path))
                
    def remove_document(self):
        selection = self.docs_listbox.curselection()
        if selection:
            index = selection[0]
            self.docs_listbox.delete(index)
            self.document_paths.pop(index)
            
    def clear_documents(self):
        self.docs_listbox.delete(0, tk.END)
        self.document_paths = []
        
    def browse_output(self):
        path = filedialog.asksaveasfilename(
            title="Save Merged Document As",
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx")]
        )
        if path:
            self.output_path_var.set(path)
            
    def merge_documents(self):
        output_path = self.output_path_var.get()
        
        if not file_exists(TEMPLATE_PATH):
            messagebox.showerror("Error", f"Template not found at: {TEMPLATE_PATH}")
            return
            
        if not self.document_paths:
            messagebox.showerror("Error", "Please add at least one document to merge.")
            return
            
        if not output_path:
            messagebox.showerror("Error", "Please specify an output document path.")
            return
            
        # Perform merge operation
        self.status_var.set("Merging documents... Please wait.")
        self.root.update()
        
        success = merge_documents(TEMPLATE_PATH, self.document_paths, output_path)
        
        if success:
            self.status_var.set(f"Documents merged successfully! Output saved to: {output_path}")
            messagebox.showinfo("Success", "Documents merged successfully!")
        else:
            self.status_var.set("Error merging documents. See console for details.")
            messagebox.showerror("Error", "Failed to merge documents. See console for details.")

def run_gui():
    """Run the GUI application"""
    # Set up the application
    root = tk.Tk()
    app = WordMergerApp(root)
    
    # Center the window on the screen
    window_width = 700
    window_height = 550
    screen_width = root.winfo_screenwidth()
    screen_height = root.winfo_screenheight()
    x_position = (screen_width - window_width) // 2
    y_position = (screen_height - window_height) // 2
    root.geometry(f"{window_width}x{window_height}+{x_position}+{y_position}")
    
    # Start the application
    root.mainloop()

def run_cli():
    """Run the command-line interface"""
    if len(sys.argv) < 3:
        print(f"Usage: python merge_word_docs.py <output.docx> <doc1.docx> [doc2.docx] [...]")
        print(f"Developed by Aaqib Jeelani - instagram.com/aaqibjeelani")
        sys.exit(1)
    
    output_path = sys.argv[1]
    document_paths = sys.argv[2:]
    
    success = merge_documents(TEMPLATE_PATH, document_paths, output_path)
    sys.exit(0 if success else 1)

# Main entry point
if __name__ == "__main__":
    # Display version and attribution
    VERSION = "1.0.0"
    print(f"Word Document Merger v{VERSION}")
    print(f"Developed by Aaqib Jeelani - instagram.com/aaqibjeelani")
    print("------------------------------------------------------")

    if len(sys.argv) > 1:
        run_cli()
    else:
        run_gui() 